"""
Generate a small sample.docx fixture next to this file. Idempotent.
Run directly or import `make_sample_docx()` from the demos.
"""
from __future__ import annotations

from pathlib import Path

HERE = Path(__file__).resolve().parent
TARGET = HERE / "sample.docx"

PARAGRAPHS = [
    ("Quarterly Status — Q3", "title"),
    ("Customer: Acme Corporation", None),
    ("Account manager: Mario Rossi (mario.rossi@acme.example)", None),
    ("Phone: +39 02 1234 5678", None),
    ("Order ref: ORD-90215", None),
    ("", None),
    ("Summary", "h1"),
    ("Acme Corporation renewed for another 12 months. Mario Rossi escalated a "
     "billing issue on order ORD-90215 last week; resolved by 2025-09-12. Next "
     "QBR scheduled with Mario Rossi and his team in Milan on 2025-10-04.", None),
    ("", None),
    ("Action items", "h1"),
    ("Send updated invoice to mario.rossi@acme.example", None),
    ("Confirm contract terms with Acme Corporation legal", None),
    ("Add ORD-90215 to the renewals dashboard", None),
]


def make_sample_docx(force: bool = False) -> Path:
    if TARGET.exists() and not force:
        return TARGET
    from docx import Document
    doc = Document()
    for text, style in PARAGRAPHS:
        if style == "title":
            doc.add_heading(text, level=0)
        elif style == "h1":
            doc.add_heading(text, level=1)
        else:
            doc.add_paragraph(text)
    doc.save(TARGET)
    return TARGET


if __name__ == "__main__":
    p = make_sample_docx(force=True)
    print(f"wrote {p}")
